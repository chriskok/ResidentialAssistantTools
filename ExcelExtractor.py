from docx import Document
from docx.shared import Inches
import pandas as pd

# READ FROM EXCEL
file_name =  "ignore/4E.xlsx"
sheet = 0
df = pd.read_excel(io=file_name, sheet_name=sheet)
cols = [3, 6, 8]
df.drop(df.columns[cols],axis=1,inplace=True)
# print(df.head(5))  # print first 5 rows of the dataframe

# WRITE TO WORD
document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Index'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Email'
for index, row in df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(index)
    row_cells[1].text = str(row.Name)
    row_cells[2].text = str(row.Email)

document.add_page_break()

document.save('ignore/demo.docx')
