from docx import Document
from docx.shared import Inches
import pandas as pd

# READ FROM EXCEL
file_name =  "ignore/4E.xlsx"
sheet = 0
df = pd.read_excel(io=file_name, sheet_name=sheet)
cols = [3, 6, 8]
df.drop(df.columns[cols],axis=1,inplace=True)

# WRITE TO WORD
document = Document()

document.add_heading('Resident Information', 0)

p = document.add_paragraph('A list of all the information on ')
p.add_run('residents of 4E.').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

document.add_heading('Table of Content', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')
# document.add_paragraph('first item in unordered list', style='List Bullet')
# document.add_paragraph('first item in ordered list', style='List Number')

table = document.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Index').bold = True
hdr_cells[1].paragraphs[0].add_run('Name').bold = True
hdr_cells[2].paragraphs[0].add_run('Email').bold = True
hdr_cells[3].paragraphs[0].add_run('Bedspace').bold = True
for index, row in df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(index)
    row_cells[1].text = str(row.Name)
    row_cells[2].text = str(row.Email)
    row_cells[3].text = str(row.Bedspace)

document.add_page_break()

for index, row in df.iterrows():
    document.add_heading(row.Name + "(" + row.Bedspace + ")", level=1)
    document.add_page_break()

document.save('ignore/demo.docx')
